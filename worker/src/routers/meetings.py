from __future__ import annotations

from typing import Any, Dict, List, Optional
from pathlib import Path
from fastapi import APIRouter, HTTPException, Query, Depends
from fastapi import Response
from fastapi.responses import PlainTextResponse, HTMLResponse

from ..db import (
    get_connection,
    new_meeting,
    delete_meeting,
    list_utterances_for_meeting,
    list_meetings,
    get_meeting_notes,
    upsert_meeting_notes,
)
from ..models.meeting import MeetingMeta, Utterance, ExportJSONResponse
from ..state import State, get_state
from ..services.notes import make_notes_from_utterances, summary_diagnostics

router = APIRouter(tags=["meetings"])


def _fetch_meeting_meta(meeting_id: int) -> MeetingMeta:
    with get_connection() as conn:
        cur = conn.cursor()
        # Prefer created_at; fall back to legacy 'created'
        try:
            cur.execute(
                "SELECT id, title, created_at FROM meetings WHERE id = ?",
                (meeting_id,),
            )
            row = cur.fetchone()
        except Exception:
            row = None
        if row is None:
            cur.execute(
                "SELECT id, title, created FROM meetings WHERE id = ?",
                (meeting_id,),
            )
            row = cur.fetchone()
        if row is None:
            raise HTTPException(status_code=404, detail=f"Meeting {meeting_id} not found")
        created_at = row[2] if len(row) > 2 else None
        return MeetingMeta(id=int(row[0]), title=row[1], created_at=created_at)  # type: ignore[arg-type]


def _fetch_meeting_utterances(meeting_id: int) -> List[Utterance]:
    items = list_utterances_for_meeting(meeting_id, limit=10_000)
    return [Utterance(**it) for it in items]


@router.post("/meeting/new")
def v1_meeting_new(title: str = "Untitled") -> Dict[str, Any]:
    mid = new_meeting(title)
    return {"ok": True, "meeting_id": mid, "title": title}


@router.get("/meetings")
def v1_meetings(limit: int = Query(200, ge=1, le=1000)) -> Dict[str, Any]:
    """List recent meetings (newest first)."""
    items = list_meetings(limit=limit)
    return {"ok": True, "items": items}


@router.delete("/meeting/{meeting_id}")
def v1_meeting_delete(
    meeting_id: int,
    cascade: bool = Query(default=True, description="If true, also delete all utterances for the meeting."),
    state: State = Depends(get_state),
) -> Dict[str, Any]:
    filenames: List[str] = []
    if cascade:
        with get_connection() as conn:
            cur = conn.cursor()
            cur.execute(
                "SELECT filename FROM utterances WHERE meeting_id = ? AND filename IS NOT NULL AND filename != ''",
                (meeting_id,),
            )
            filenames = [r[0] for r in cur.fetchall() if r and r[0]]

    try:
        deleted_meeting, deleted_utt = delete_meeting(meeting_id, cascade=cascade)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    if deleted_meeting == 0:
        raise HTTPException(status_code=404, detail=f"Meeting {meeting_id} not found")

    deleted_audio = 0
    if cascade and deleted_meeting and filenames:
        base_dir = Path(state.tmp_dir).resolve()
        for name in filenames:
            try:
                p = Path(str(name))
                if p.is_absolute():
                    candidate = p
                else:
                    candidate = base_dir / p.name
                candidate = candidate.resolve()
                try:
                    candidate.relative_to(base_dir)
                except Exception:
                    continue
                if candidate.is_file():
                    candidate.unlink()
                    deleted_audio += 1
            except Exception:
                continue

    deleted_exports = 0
    if cascade and deleted_meeting:
        exports_dir = Path(__file__).resolve().parent.parent / "exports"
        if exports_dir.is_dir():
            patterns = [
                f"meeting-{meeting_id}-*.md",
                f"meeting-{meeting_id}-*.docx",
            ]
            for pattern in patterns:
                for path in exports_dir.glob(pattern):
                    try:
                        if path.is_file():
                            path.unlink()
                            deleted_exports += 1
                    except Exception:
                        continue

    return {
        "ok": True,
        "meeting_id": meeting_id,
        "deleted_meeting": deleted_meeting,
        "deleted_utterances": deleted_utt,
        "deleted_audio_files": deleted_audio,
        "deleted_export_files": deleted_exports,
    }


@router.post("/meeting/start/{meeting_id}")
def v1_meeting_start(meeting_id: int, state: State = Depends(get_state)) -> Dict[str, Any]:
    # Validate meeting exists
    with get_connection() as conn:
        cur = conn.cursor()
        cur.execute("SELECT 1 FROM meetings WHERE id = ?", (meeting_id,))
        if cur.fetchone() is None:
            raise HTTPException(status_code=404, detail="Meeting not found")
    state.current_meeting_id = meeting_id
    return {"ok": True, "id": meeting_id}


@router.post("/meeting/stop")
def v1_meeting_stop(state: State = Depends(get_state)) -> Dict[str, Any]:
    state.current_meeting_id = None
    return {"ok": True}


@router.get("/meeting/active")
def v1_meeting_active(state: State = Depends(get_state)) -> Dict[str, Any]:
    return {"id": state.current_meeting_id}


@router.get("/meeting/{meeting_id}/export.json", response_model=ExportJSONResponse)
def v1_export_meeting_json(
    meeting_id: int,
    speakers: str = Query("0", description="'0'|'heuristic'|'auto' (v1 currently supports '0' only)"),
    max_speakers: Optional[int] = Query(None, ge=1, le=20),
) -> ExportJSONResponse:
    if speakers not in {"0", "heuristic", "auto"}:
        raise HTTPException(status_code=400, detail="Invalid speakers mode")
    if speakers != "0":
        # Placeholder: diarization will be added in a later step under services/
        # For now, keep parity with legacy endpoints by returning without labels.
        pass
    meta = _fetch_meeting_meta(meeting_id)
    utterances = _fetch_meeting_utterances(meeting_id)
    return ExportJSONResponse(meeting=meta, utterances=utterances)


@router.get("/meeting/{meeting_id}/export.md", response_class=PlainTextResponse)
def v1_export_meeting_markdown(
    meeting_id: int,
    speakers: str = Query("0", description="'0'|'heuristic'|'auto' (v1 currently supports '0' only)"),
    max_speakers: Optional[int] = Query(None, ge=1, le=20),
) -> str:
    if speakers not in {"0", "heuristic", "auto"}:
        raise HTTPException(status_code=400, detail="Invalid speakers mode")
    # Placeholder: diarization to be introduced later
    meta = _fetch_meeting_meta(meeting_id)
    utterances = _fetch_meeting_utterances(meeting_id)

    header_lines: List[str] = [f"# {meta.title or 'Untitled'}"]
    if meta.created_at:
        header_lines.append(f"_Created: {meta.created_at}_")
    header_lines.append("")

    lines: List[str] = []
    for u in utterances:
        ts = u.ts_iso or ""
        text = u.text or ""
        if u.speaker:
            lines.append(f"- **{u.speaker}** [{ts}] {text}")
        else:
            lines.append(f"- [{ts}] {text}")
    return "\n".join(header_lines + lines) + "\n"


@router.get("/meeting/{meeting_id}/utterances")
def v1_list_utterances(
    meeting_id: int,
    since_id: int = Query(0, ge=0),
    limit: int = Query(200, ge=1, le=1000),
):
    with get_connection() as conn:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, meeting_id, ts_iso, start_ms, end_ms, text, confidence, filename
            FROM utterances
            WHERE meeting_id = ? AND id > ?
            ORDER BY id ASC
            LIMIT ?
            """,
            (meeting_id, since_id, limit),
        )
        rows = cur.fetchall()
    items: List[Dict[str, Any]] = []
    for r in rows:
        items.append(
            {
                "id": r[0],
                "meeting_id": r[1],
                "ts_iso": r[2],
                "start_ms": r[3],
                "end_ms": r[4],
                "text": r[5],
                "confidence": r[6],
                "filename": r[7],
            }
        )
    next_since = items[-1]["id"] if items else since_id
    return {"ok": True, "items": items, "next_since_id": next_since}


@router.get("/summary_diagnostics")
def v1_summary_diagnostics() -> Dict[str, Any]:
    """Expose summary provider diagnostics for debugging configuration and connectivity."""
    info = summary_diagnostics()
    return {"ok": True, **info}


@router.get("/meeting/{meeting_id}/summary", response_class=HTMLResponse)
def v1_meeting_summary(
    meeting_id: int,
    strict: bool = Query(False),
    regen: bool = Query(False),
    response: Response = None,
) -> str:
    meta = _fetch_meeting_meta(meeting_id)
    cached = None if regen else get_meeting_notes(meeting_id, strict=strict)
    if cached is not None:
        notes_md = cached["content_md"]
        provider = cached.get("provider")
        model = cached.get("model")
        created_at = cached.get("created_at")
    else:
        utterances = _fetch_meeting_utterances(meeting_id)
        notes_md = make_notes_from_utterances([u.dict() for u in utterances], strict=strict)
        diag = summary_diagnostics()
        provider = diag.get("provider")
        model = diag.get("model")
        upsert_meeting_notes(meeting_id, notes_md, strict=strict, provider=provider, model=model)

    # Simple HTML wrapper; render markdown-like content minimally
    def md_to_html(md: str) -> str:
        # Minimal conversion for headers and bullets
        lines = []
        for line in md.splitlines():
            if line.startswith("## "):
                lines.append(f"<h2>{line[3:].strip()}</h2>")
            elif line.startswith("- "):
                # Start/close lists automatically
                if not lines or not lines[-1].startswith("<ul>"):
                    lines.append("<ul>")
                lines.append(f"<li>{line[2:].strip()}</li>")
            elif line.strip() == "":
                if lines and lines[-1] == "<ul>":
                    # do nothing for empty lines at start of list
                    pass
                else:
                    lines.append("<p></p>")
            else:
                lines.append(f"<p>{line}</p>")
        # Close dangling list
        out: list[str] = []
        open_list = False
        for l in lines:
            if l == "<ul>":
                if not open_list:
                    out.append("<ul>")
                    open_list = True
            else:
                if open_list and not l.startswith("<li>"):
                    out.append("</ul>")
                    open_list = False
                out.append(l)
        if open_list:
            out.append("</ul>")
        return "\n".join(out)

    body = md_to_html(notes_md)
    title = meta.title or f"Meeting {meta.id}"
    created = f"<div class=\"muted\">Created: {meta.created_at}</div>" if meta.created_at else ""
    badge = ""
    try:
        if provider:
            badge = f"<div class=\"muted\" style=\"margin:6px 0\">Generated by: {provider} {model or ''}</div>"
    except Exception:
        pass
    html = f"""
    <!doctype html>
    <html>
    <head>
      <meta charset=\"utf-8\" />
      <title>Summary – {title}</title>
      <style>
        body{{font:14px/1.5 -apple-system,Segoe UI,Roboto,Inter,Helvetica,Arial,sans-serif; margin:24px; color:#111}}
        .muted{{color:#666; margin:4px 0 16px}}
        h1{{font-size:18px; margin:0 0 4px}}
        h2{{font-size:16px; margin:18px 0 6px}}
        ul{{margin:6px 0 10px 20px}}
        li{{margin:4px 0}}
        p{{margin:6px 0}}
      </style>
    </head>
    <body>
      <h1>{title} – Notes</h1>
      {created}
      {badge}
      {body}
    </body>
    </html>
    """
    # Prevent browser/proxy caching
    try:
        if response is not None:
            response.headers["Cache-Control"] = "no-store"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
    except Exception:
        pass
    return html


@router.get("/meeting/{meeting_id}/summary.md", response_class=PlainTextResponse)
def v1_meeting_summary_markdown(
    meeting_id: int,
    strict: bool = Query(False),
    regen: bool = Query(False),
    response: Response = None,
) -> str:
    meta = _fetch_meeting_meta(meeting_id)
    cached = None if regen else get_meeting_notes(meeting_id, strict=strict)
    if cached is not None:
        body_md = cached["content_md"]
        provider = cached.get("provider")
        model = cached.get("model")
    else:
        utterances = _fetch_meeting_utterances(meeting_id)
        body_md = make_notes_from_utterances([u.dict() for u in utterances], strict=strict)
        diag = summary_diagnostics()
        provider = diag.get("provider")
        model = diag.get("model")
        upsert_meeting_notes(meeting_id, body_md, strict=strict, provider=provider, model=model)

    title = meta.title or f"Meeting {meta.id}"
    header_lines = [f"# {title} – Notes"]
    if meta.created_at:
        header_lines.append(f"_Created: {meta.created_at}_")
    header_lines.append("")
    if provider:
        header_lines.append(f"_Generated by: {provider} {model or ''}_")
        header_lines.append("")
    out = "\n".join(header_lines + [body_md.strip(), ""])  # trailing newline
    # Prevent browser/proxy caching
    try:
        if response is not None:
            response.headers["Cache-Control"] = "no-store"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
    except Exception:
        pass
    return out


@router.post("/meeting/{meeting_id}/export_notes")
def v1_meeting_export_notes(
    meeting_id: int,
    strict: bool = Query(False),
    regen: bool = Query(False),
) -> Dict[str, Any]:
    """Save the (cached or freshly generated) notes to worker/exports as Markdown.

    Returns { ok, path, filename } with absolute path for convenience.
    """
    # Prepare content (reuse cache unless regen)
    meta = _fetch_meeting_meta(meeting_id)
    cached = None if regen else get_meeting_notes(meeting_id, strict=strict)
    if cached is not None:
        body_md = cached["content_md"]
        provider = cached.get("provider")
        model = cached.get("model")
    else:
        utterances = _fetch_meeting_utterances(meeting_id)
        body_md = make_notes_from_utterances([u.dict() for u in utterances], strict=strict)
        diag = summary_diagnostics()
        provider = diag.get("provider")
        model = diag.get("model")
        upsert_meeting_notes(meeting_id, body_md, strict=strict, provider=provider, model=model)

    # Compose Markdown (header + body)
    title = meta.title or f"Meeting {meta.id}"
    header_lines = [f"# {title} – Notes"]
    if meta.created_at:
        header_lines.append(f"_Created: {meta.created_at}_")
    if provider:
        header_lines.append(f"_Generated by: {provider} {model or ''}_")
    header_lines.append("")
    content = "\n".join(header_lines + [body_md.strip(), ""])  # trailing newline

    # Determine exports directory (worker/exports)
    from pathlib import Path
    import re
    base_dir = Path(__file__).resolve().parent.parent  # worker/
    exports_dir = base_dir / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)

    # Sanitize filename
    safe_title = re.sub(r"[^a-zA-Z0-9-_]+", "-", (title or "meeting")).strip("-")
    suffix = "-strict" if strict else ""
    filename = f"meeting-{meta.id}-{safe_title}{suffix}.md"
    path = exports_dir / filename

    path.write_text(content, encoding="utf-8")
    return {"ok": True, "path": str(path), "filename": filename}


@router.post("/meeting/{meeting_id}/export_notes_docx")
def v1_meeting_export_notes_docx(
    meeting_id: int,
    strict: bool = Query(False),
    regen: bool = Query(False),
) -> Dict[str, Any]:
    """Save notes to worker/exports as a .docx file.

    Requires python-docx. If not installed, returns an error with guidance.
    """
    # Lazy import to allow worker to run without the dependency
    try:
        from docx import Document  # type: ignore
    except Exception:
        return {
            "ok": False,
            "error": "python-docx not installed",
            "install": "pip install python-docx",
        }

    meta = _fetch_meeting_meta(meeting_id)
    cached = None if regen else get_meeting_notes(meeting_id, strict=strict)
    if cached is not None:
        body_md = cached["content_md"]
        provider = cached.get("provider")
        model = cached.get("model")
    else:
        utterances = _fetch_meeting_utterances(meeting_id)
        body_md = make_notes_from_utterances([u.dict() for u in utterances], strict=strict)
        diag = summary_diagnostics()
        provider = diag.get("provider")
        model = diag.get("model")
        upsert_meeting_notes(meeting_id, body_md, strict=strict, provider=provider, model=model)

    # Create the document
    doc = Document()
    title = meta.title or f"Meeting {meta.id}"
    doc.add_heading(f"{title} – Notes", level=1)
    if meta.created_at:
        doc.add_paragraph(f"Created: {meta.created_at}")
    if provider:
        doc.add_paragraph(f"Generated by: {provider} {model or ''}")

    # Very simple Markdown-ish to docx: map headings and bullets
    for line in (body_md or "").splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    # Write file
    from pathlib import Path
    import re
    base_dir = Path(__file__).resolve().parent.parent  # worker/
    exports_dir = base_dir / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r"[^a-zA-Z0-9-_]+", "-", (title or "meeting")).strip("-")
    suffix = "-strict" if strict else ""
    filename = f"meeting-{meta.id}-{safe_title}{suffix}.docx"
    path = exports_dir / filename
    doc.save(str(path))
    return {"ok": True, "path": str(path), "filename": filename}
